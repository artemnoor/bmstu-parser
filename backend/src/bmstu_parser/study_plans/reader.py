from __future__ import annotations

import hashlib
import re
import shutil
import subprocess
import tempfile
import xml.etree.ElementTree as ET
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from ..domain.ids import stable_id
from .ids import row_id, table_id
from .models import ExtractedDocument, StudyPlanReference
from .readers import DocumentReader, get_reader_backend


def utc_now() -> str:
    return datetime.now(UTC).isoformat()


def _local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def _hash_file(path: Path) -> tuple[int, str, bytes]:
    digest = hashlib.sha256()
    head = b""
    size = 0
    with path.open("rb") as stream:
        while chunk := stream.read(1024 * 1024):
            if not head:
                head = chunk[:64]
            digest.update(chunk)
            size += len(chunk)
    return size, digest.hexdigest(), head


def _int_or_none(value: Any) -> int | None:
    try:
        return int(value) if value not in (None, "") else None
    except (TypeError, ValueError):
        return None


def _source_kind(path: Path, head: bytes) -> str:
    if head.startswith(b"%PDF-"):
        return "pdf"
    if head.startswith(b"PK\x03\x04") and path.suffix.lower() == ".docx":
        return "docx"
    prefix = head.lstrip().lower()
    if prefix.startswith(b"<!doctype") or b"<html" in prefix[:512]:
        return "html"
    return "unknown"


def _xml_pages(xml_bytes: bytes) -> list[dict[str, Any]]:
    root = ET.fromstring(xml_bytes)
    pages: list[dict[str, Any]] = []
    for page_number, page_element in enumerate(
        (element for element in root.iter() if _local_name(element.tag) == "page"),
        start=1,
    ):
        words: list[dict[str, Any]] = []
        lines: list[dict[str, Any]] = []
        for line_element in (
            element
            for element in page_element.iter()
            if _local_name(element.tag) == "line"
        ):
            line_words: list[dict[str, Any]] = []
            for word_element in (
                element
                for element in line_element
                if _local_name(element.tag) == "word"
            ):
                text = "".join(word_element.itertext())
                if not text:
                    continue
                attrs = word_element.attrib
                word = {
                    "id": f"word-{page_number}-{len(words)}",
                    "text": text,
                    "x0": float(attrs.get("xMin", 0)),
                    "top": float(attrs.get("yMin", 0)),
                    "x1": float(attrs.get("xMax", 0)),
                    "bottom": float(attrs.get("yMax", 0)),
                }
                words.append(word)
                line_words.append(word)
            if line_words:
                lines.append(
                    {
                        "line_index": len(lines),
                        "text": " ".join(word["text"] for word in line_words),
                        "top": min(word["top"] for word in line_words),
                        "bottom": max(word["bottom"] for word in line_words),
                        "word_ids": [word["id"] for word in line_words],
                    }
                )
        pages.append(
            {
                "page_number": page_number,
                "width": float(page_element.attrib.get("width", 0)),
                "height": float(page_element.attrib.get("height", 0)),
                "words": words,
                "lines": lines,
                "table_ids": [],
            }
        )
    return pages


def _words_in_cell(
    words: list[dict[str, Any]], cell: tuple[float, float, float, float]
) -> list[dict[str, Any]]:
    x0, top, x1, bottom = cell
    selected = []
    for word in words:
        center_x = (word["x0"] + word["x1"]) / 2
        center_y = (word["top"] + word["bottom"]) / 2
        if x0 - 0.6 <= center_x <= x1 + 0.6 and top - 0.6 <= center_y <= bottom + 0.6:
            selected.append(word)
    return sorted(selected, key=lambda item: (item["top"], item["x0"]))


def _cell_text(words: list[dict[str, Any]]) -> str:
    if not words:
        return ""
    lines: list[list[dict[str, Any]]] = []
    for word in words:
        if not lines or abs(word["top"] - lines[-1][0]["top"]) > 1.6:
            lines.append([word])
        else:
            lines[-1].append(word)
    return "\n".join(" ".join(word["text"] for word in line) for line in lines)


def _classify_table(table: Any) -> str:
    columns = len(table.columns)
    x0, _top, _x1, _bottom = table.bbox
    if columns >= 60:
        return "curriculum"
    if columns >= 25:
        return "calendar_schedule"
    if x0 >= 700 or columns <= 10:
        return "time_budget_summary"
    return "other"


def _extract_pdf(
    path: Path, document_id: str
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], str, list[str]]:
    pdftotext = shutil.which("pdftotext")
    if not pdftotext:
        raise RuntimeError("Не найден pdftotext из Poppler")
    temp_dir: tempfile.TemporaryDirectory[str] | None = None
    source_path = path
    if any(ord(character) > 127 for character in str(path)):
        temp_dir = tempfile.TemporaryDirectory(prefix="bmstu_pdf_")
        source_path = Path(temp_dir.name) / "source.pdf"
        shutil.copyfile(path, source_path)
    try:
        bbox_process = subprocess.run(
            [pdftotext, "-bbox-layout", "-enc", "UTF-8", str(source_path), "-"],
            check=False,
            capture_output=True,
        )
        if bbox_process.returncode != 0:
            raise RuntimeError(
                (bbox_process.stderr or b"pdftotext failed").decode(
                    "utf-8", errors="replace"
                )
            )
        layout_process = subprocess.run(
            [pdftotext, "-layout", "-enc", "UTF-8", str(source_path), "-"],
            check=False,
            capture_output=True,
        )
        if layout_process.returncode != 0:
            raise RuntimeError(
                (layout_process.stderr or b"pdftotext failed").decode(
                    "utf-8", errors="replace"
                )
            )
        pages = _xml_pages(bbox_process.stdout)
        tables: list[dict[str, Any]] = []
        warnings: list[str] = []

        try:
            import pdfplumber

            settings = {
                "vertical_strategy": "lines",
                "horizontal_strategy": "lines",
                "snap_tolerance": 3,
                "join_tolerance": 3,
                "intersection_tolerance": 3,
            }
            with pdfplumber.open(str(source_path)) as pdf:
                if len(pdf.pages) != len(pages):
                    warnings.append(
                        f"Количество страниц отличается: bbox={len(pages)}, pdfplumber={len(pdf.pages)}"
                    )
                for page_number, page in enumerate(pdf.pages, start=1):
                    if page_number > len(pages):
                        break
                    try:
                        detected = page.find_tables(table_settings=settings)
                    # pdfplumber exposes backend-specific exceptions for malformed tables.
                    except Exception as exc:  # noqa: BLE001
                        warnings.append(
                            f"Страница {page_number}: ошибка поиска таблиц: {exc}"
                        )
                        detected = []
                    page_words = pages[page_number - 1]["words"]
                    for table_index, detected_table in enumerate(detected):
                        identifier = table_id(document_id, page_number, table_index)
                        section = _classify_table(detected_table)
                        cell_rows: list[list[dict[str, Any]]] = []
                        for row_index, row in enumerate(detected_table.rows):
                            cell_row: list[dict[str, Any]] = []
                            for column_index, cell in enumerate(row.cells):
                                cell_identifier = stable_id(
                                    "study-plan-cell",
                                    identifier,
                                    row_index,
                                    column_index,
                                )
                                if cell is None:
                                    cell_row.append(
                                        {
                                            "id": cell_identifier,
                                            "table_id": identifier,
                                            "row_index": row_index,
                                            "column_index": column_index,
                                            "text": "",
                                            "bbox": None,
                                            "word_ids": [],
                                            "cell_kind": "merged_placeholder",
                                        }
                                    )
                                    continue
                                selected_words = _words_in_cell(page_words, cell)
                                cell_row.append(
                                    {
                                        "id": cell_identifier,
                                        "table_id": identifier,
                                        "row_index": row_index,
                                        "column_index": column_index,
                                        "text": _cell_text(selected_words),
                                        "bbox": {
                                            "x0": cell[0],
                                            "top": cell[1],
                                            "x1": cell[2],
                                            "bottom": cell[3],
                                        },
                                        "word_ids": [
                                            word["id"] for word in selected_words
                                        ],
                                        "cell_kind": "cell",
                                    }
                                )
                            cell_rows.append(cell_row)
                        record = {
                            "id": identifier,
                            "document_id": document_id,
                            "page_number": page_number,
                            "table_index": table_index,
                            "section": section,
                            "bbox": {
                                "x0": detected_table.bbox[0],
                                "top": detected_table.bbox[1],
                                "x1": detected_table.bbox[2],
                                "bottom": detected_table.bbox[3],
                            },
                            "row_count": len(cell_rows),
                            "column_count": max(
                                (len(row) for row in cell_rows), default=0
                            ),
                            "rows": cell_rows,
                            "extraction_method": "pdfplumber-grid+pdftotext-bbox",
                        }
                        tables.append(record)
                        pages[page_number - 1]["table_ids"].append(identifier)
        except (
            ImportError
        ) as exc:  # pragma: no cover - dependency is declared in pyproject
            raise RuntimeError(f"Не установлен pdfplumber: {exc}") from exc

        if not tables:
            warnings.append("На страницах не обнаружены таблицы по линиям")
        return (
            pages,
            tables,
            layout_process.stdout.decode("utf-8", errors="replace"),
            warnings,
        )
    finally:
        if temp_dir is not None:
            temp_dir.cleanup()


def _extract_docx(
    path: Path, document_id: str
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], str, list[str]]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency is declared in pyproject
        raise RuntimeError(f"Не установлен python-docx: {exc}") from exc
    document = Document(str(path))
    paragraphs = [
        {"paragraph_index": index, "text": paragraph.text}
        for index, paragraph in enumerate(document.paragraphs)
    ]
    tables: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables):
        identifier = table_id(document_id, 1, table_index)
        rows: list[list[dict[str, Any]]] = []
        for row_index, row in enumerate(table.rows):
            cell_row = []
            for column_index, cell in enumerate(row.cells):
                cell_row.append(
                    {
                        "id": stable_id(
                            "study-plan-cell", identifier, row_index, column_index
                        ),
                        "table_id": identifier,
                        "row_index": row_index,
                        "column_index": column_index,
                        "text": cell.text,
                        "bbox": None,
                        "word_ids": [],
                        "cell_kind": "office_cell",
                    }
                )
            rows.append(cell_row)
        tables.append(
            {
                "id": identifier,
                "document_id": document_id,
                "page_number": 1,
                "table_index": table_index,
                "section": "office_table",
                "bbox": None,
                "row_count": len(rows),
                "column_count": max((len(row) for row in rows), default=0),
                "rows": rows,
                "extraction_method": "python-docx",
            }
        )
    return (
        [
            {
                "page_number": 1,
                "width": None,
                "height": None,
                "words": [],
                "lines": paragraphs,
                "table_ids": [table["id"] for table in tables],
            }
        ],
        tables,
        "\n".join(item["text"] for item in paragraphs),
        ["В DOCX не обнаружены таблицы"] if not tables else [],
    )


def extract_document(
    reference: StudyPlanReference,
    root: Path,
    reader_backend: str | DocumentReader = "native",
) -> dict[str, Any]:
    path = root / Path(reference.local_path.replace("\\", "/"))
    source_size, source_sha256, head = (
        _hash_file(path) if path.exists() else (0, "", b"")
    )
    # Slots dataclasses do not expose __dict__; keep the explicit representation deterministic.
    source_references = [
        {
            "document_id": reference.document_id,
            "local_path": reference.local_path,
            "major_id": reference.major_id,
            "major_slug": reference.major_slug,
            "major_code": reference.major_code,
            "major_name": reference.major_name,
            "program_id": reference.program_id,
            "program_code": reference.program_code,
            "program_name": reference.program_name,
            "plan_url": reference.plan_url,
            "plan_status": reference.plan_status,
            "source_url": reference.source_url,
            "resolved_url": reference.resolved_url,
            "expected_size": reference.expected_size,
            "expected_sha256": reference.expected_sha256,
            "expected_mime_type": reference.expected_mime_type,
        }
    ]
    base = {
        "document": ExtractedDocument(
            document_id=reference.document_id,
            local_path=reference.local_path,
            absolute_path=str(path.resolve()),
            kind="missing" if not path.exists() else _source_kind(path, head),
            status="missing" if not path.exists() else "pending",
            source_references=source_references,
            source_url=reference.source_url,
            resolved_url=reference.resolved_url,
            source_size=source_size,
            source_sha256=source_sha256,
            expected_size=reference.expected_size,
            expected_sha256=reference.expected_sha256,
            expected_mime_type=reference.expected_mime_type,
            extracted_at_utc=utc_now(),
        ).to_dict(),
        "pages": [],
        "tables": [],
        "layout_text": "",
    }
    if not path.exists():
        base["document"]["errors"] = [f"Файл не найден: {path}"]
        return base
    try:
        kind = base["document"]["kind"]
        if kind not in {"pdf", "docx"}:
            base["document"]["status"] = "invalid_source"
            base["document"]["errors"] = [f"Ожидался PDF/DOCX, получен {kind}"]
            return base
        backend = get_reader_backend(reader_backend)
        pages, tables, layout_text, warnings = backend.extract(
            path, reference.document_id
        )
        base["pages"] = pages
        base["tables"] = tables
        base["layout_text"] = layout_text
        base["document"]["extraction_backend"] = backend.name
        base["document"]["status"] = "ok" if tables else "ok_no_tables"
        base["document"]["page_count"] = len(pages)
        base["document"]["paragraph_count"] = sum(
            len(page.get("lines", [])) for page in pages
        )
        base["document"]["table_count"] = len(tables)
        base["document"]["row_count"] = sum(table["row_count"] for table in tables)
        base["document"]["cell_count"] = sum(
            len(row) for table in tables for row in table.get("rows", [])
        )
        base["document"]["warnings"] = warnings
    # Preserve an auditable failed document rather than dropping it.
    except Exception as exc:  # noqa: BLE001
        base["document"]["status"] = "error"
        base["document"]["errors"] = [f"{type(exc).__name__}: {exc}"]
    return base


def curriculum_row_record(
    table: dict[str, Any], row_index: int, cells: list[dict[str, Any]]
) -> dict[str, Any]:
    values = [cell.get("text", "").strip() for cell in cells]
    first = values[0] if values else ""
    second = values[1] if len(values) > 1 else ""
    code = first if re.fullmatch(r"(?:[А-ЯA-Za-z]\d(?:[.\d]*)?|\d+)", first) else ""
    if code:
        role = "discipline" if code.isdigit() else "section"
    elif any(marker in second.lower() for marker in ("часть", "дисциплины", "модул")):
        role = "summary"
    elif row_index < 5:
        role = "header"
    else:
        role = "continuation_or_summary"
    return {
        "id": row_id(table["id"], row_index),
        "table_id": table["id"],
        "document_id": table["document_id"],
        "page_number": table["page_number"],
        "section": table["section"],
        "row_index": row_index,
        "row_role": role,
        "code": code,
        "name": second,
        "cells": cells,
    }
